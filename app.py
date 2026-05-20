import sqlite3
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file
import random
import os  # <-- Make sure this import is present
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.pdfgen import canvas
app = Flask(__name__)
app.secret_key = 'your_secret_key'  # Necessary for flash messages to work

# Path to the ssl certificate and key
ssl_cert = 'ssl/server.crt'
ssl_key = 'ssl/server.key'

# Path to the file containing usernames
usernames_file = 'usernames.txt'
DB_FILE = 'votestack2.db'


def get_db_connection():
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    conn = get_db_connection()
    cursor = conn.cursor()

    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        phone_number TEXT UNIQUE NOT NULL
    )
    ''')

    cursor.execute('''
    CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT
    )
    ''')

    cursor.execute('''
    CREATE TABLE IF NOT EXISTS choices (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        role TEXT NOT NULL,
        choice TEXT NOT NULL,
        UNIQUE(role, choice)
    )
    ''')

    cursor.execute('''
    CREATE TABLE IF NOT EXISTS votes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT NOT NULL,
        role TEXT NOT NULL,
        choice TEXT NOT NULL,
        created_at DATETIME DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(username, role)
    )
    ''')

    conn.commit()
    conn.close()


def get_current_role():
    conn = get_db_connection()
    row = conn.execute("SELECT value FROM settings WHERE key='current_role'").fetchone()
    conn.close()
    return row['value'] if row else None


def set_current_role(role):
    conn = get_db_connection()
    conn.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ('current_role', role))
    conn.commit()
    conn.close()


def get_choices_for_role(role):
    if not role:
        return []

    conn = get_db_connection()
    rows = conn.execute("SELECT choice FROM choices WHERE role=? ORDER BY id", (role,)).fetchall()
    conn.close()
    return [row['choice'] for row in rows]


def save_choices_for_role(role, choices):
    conn = get_db_connection()
    conn.execute("DELETE FROM choices WHERE role=?", (role,))
    conn.executemany(
        "INSERT OR IGNORE INTO choices (role, choice) VALUES (?, ?)",
        [(role, choice) for choice in choices if choice.strip()]
    )
    conn.commit()
    conn.close()


def has_user_voted(username, role):
    conn = get_db_connection()
    row = conn.execute(
        "SELECT 1 FROM votes WHERE username=? AND role=?",
        (username, role)
    ).fetchone()
    conn.close()
    return row is not None


def record_vote(username, role, choice):
    conn = get_db_connection()
    conn.execute(
        "INSERT INTO votes (username, role, choice) VALUES (?, ?, ?)",
        (username, role, choice)
    )
    conn.commit()
    conn.close()



@app.route('/login', methods=['POST'])
def login():

    session.clear()

    phone = request.form.get('username')  # still using same input field

    if not phone:
        flash("Please enter a phone number.", "danger")
        return redirect(url_for('index'))

    phone = ''.join(phone.split())  # remove spaces

    # admin bypass stays the same (optional)
    if phone == 'admin':
        session['username'] = phone
        return redirect(url_for('admin_dashboard'))

    try:
        conn = sqlite3.connect('votestack2.db')
        cursor = conn.cursor()

        cursor.execute(
            "SELECT * FROM users WHERE phone_number=?",
            (phone,)
        )

        user = cursor.fetchone()
        conn.close()

        if user:
            session['username'] = phone
            flash(f"Welcome {phone}", "success")
            return redirect(url_for('vote'))

        else:
            flash("Phone number not authorized", "danger")
            return redirect(url_for('index'))

    except Exception as e:
        flash(f"Database error: {e}", "danger")
        return redirect(url_for('index'))


def load_role():
    return get_current_role()

@app.route('/view_role')
def view_role():
    role = load_role()
    return render_template('view_role.html', role=role)

@app.route('/', methods=['GET'])
def index():
    return render_template('login.html')  # Show the login form

@app.route('/vote')
def vote():
    role = load_role()
    choices = get_choices_for_role(role)

    if not role or not choices:
        return "No role or choices available to vote on."

    return render_template('vote.html', role=role, choices=choices)


@app.route('/submit_vote', methods=['POST'])
def submit_vote():
    if 'username' not in session:
        return redirect(url_for('index'))  # Redirect to login if not logged in

    role = load_role()
    if not role:
        flash("No voting role is configured.", "danger")
        return redirect(url_for('vote'))

    choice = request.form.get('choice')  # Get the selected choice from the form
    if not choice:
        flash("Please select a choice to vote for.", "danger")
        return redirect(url_for('vote'))

    username = session['username']

    try:
        if has_user_voted(username, role):
            flash(f"You have already voted for the role '{role}', {username}. You can only vote once per role.", "danger")
            return redirect(url_for('vote'))

        record_vote(username, role, choice)
        flash(f"Your vote for {choice} has been recorded for role '{role}'!", "success")
        return redirect(url_for('vote'))

    except Exception as e:
        flash(f"An error occurred: {str(e)}", "danger")
        return redirect(url_for('vote'))


@app.route('/api/current_role', methods=['GET'])
def get_current_role_api():
    """API endpoint to get the current role."""
    role = load_role()
    return jsonify({"role": role}), 200

@app.route('/api/submit_vote', methods=['POST'])
def submit_vote_api():
    """API endpoint for submitting votes."""
    data = request.get_json()

    if 'username' not in session:
        return jsonify({"error": "Please log in first."}), 401

    username = session['username']
    choice = data.get('choice')
    role = load_role()

    if not role:
        return jsonify({"error": "No voting role is configured."}), 400

    if not choice:
        return jsonify({"error": "Please select a choice to vote for."}), 400

    try:
        if has_user_voted(username, role):
            return jsonify({"error": "You have already voted."}), 400

        record_vote(username, role, choice)
        return jsonify({"message": f"Your vote for {choice} has been recorded for role '{role}'!"}), 200

    except Exception as e:
        return jsonify({"error": f"An error occurred: {str(e)}"}), 500
# Route to view the current choices
@app.route('/view_choices', methods=['GET'])
def view_choices():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))

    role = load_role()
    choices = get_choices_for_role(role)

    if not choices:
        flash("No choices available.", "warning")

    return render_template('view_choices.html', choices=choices)

# Route to update the choices (allow admin to input new choices)
@app.route('/update_choices', methods=['POST'])
def update_choices():
    if request.method == 'POST':
        role = load_role()
        if not role:
            flash("No current role configured.", category='danger')
            return redirect(url_for('admin_dashboard'))

        new_choices = request.form.get('choices').splitlines()
        try:
            save_choices_for_role(role, new_choices)
            flash("Choices have been updated successfully!", category='success')
        except Exception as e:
            flash(f"An error occurred while updating choices: {e}", category='danger')

        return redirect(url_for('admin_dashboard'))

@app.route('/admin_dashboard', methods=['GET', 'POST'])
def admin_dashboard():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))  # Redirect if the user is not admin

    if request.method == 'POST':
        # New feature to add voters
        num_voters = request.form.get('num_voters')

        if not num_voters or not num_voters.isdigit() or int(num_voters) <= 0:
            flash("Please enter a valid number of voters.", "danger")
            return redirect(url_for('admin_dashboard'))

        num_voters = int(num_voters)

        # Generate random 4-digit numbers for voters
        generated_usernames = [str(random.randint(1000, 9999)) for _ in range(num_voters)]

        try:
            # Update the 'usernames.txt' file with the generated usernames
            with open(usernames_file, 'a') as file:
                for username in generated_usernames:
                    file.write(f"{username}\n")

            flash(f"{num_voters} new voters have been added.", "success")
        except Exception as e:
            flash(f"An error occurred while updating usernames: {str(e)}", "danger")

    role = load_role()
    return render_template('admin_dashboard.html', role=role)


@app.route('/generate_tally', methods=['GET', 'POST'])
def generate_tally():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))

    role = load_role()
    if not role:
        flash("No role is currently set.", "warning")
        return redirect(url_for('admin_dashboard'))

    conn = get_db_connection()
    rows = conn.execute(
        "SELECT choice, COUNT(*) AS count FROM votes WHERE role=? GROUP BY choice",
        (role,)
    ).fetchall()
    conn.close()

    tally = {row['choice']: row['count'] for row in rows}
    return render_template('tally.html', role=role, tally=tally)


@app.route('/enter_voters', methods=['POST'])
def enter_voters():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))  # Ensure only admin can access this route

    num_voters = request.form.get('num_voters')  # Get the number of voters from the form

    if not num_voters or not num_voters.isdigit() or int(num_voters) <= 0:
        flash("Please enter a valid number of voters.", "danger")
        return redirect(url_for('admin_dashboard'))  # Redirect to admin dashboard if invalid number

    num_voters = int(num_voters)

    # Generate random 4-digit numbers for voters
    generated_usernames = [str(random.randint(1000, 9999)) for _ in range(num_voters)]

    try:
        # Read existing usernames from the 'usernames.txt' file to avoid duplicates
        with open(usernames_file, 'r') as file:
            existing_usernames = set(file.read().splitlines())  # Use a set to avoid duplicates

        # Filter out usernames that already exist in 'usernames.txt'
        new_usernames = [username for username in generated_usernames if username not in existing_usernames]

        if not new_usernames:
            flash("All generated usernames already exist. No new usernames were added.", "warning")
            return redirect(url_for('admin_dashboard'))  # Redirect if no new usernames were generated

        # Update the 'usernames.txt' file with the new unique usernames
        with open(usernames_file, 'a') as file:
            for username in new_usernames:
                file.write(f"{username}\n")

        # Now, generate the PDF with the updated list of usernames
        all_usernames = list(existing_usernames) + new_usernames
        generate_pdf(all_usernames)

        flash(f"{len(new_usernames)} new voters have been added and a PDF has been generated.", "success")
        return redirect(url_for('admin_dashboard'))  # Redirect back to admin dashboard

    except Exception as e:
        flash(f"An error occurred while updating usernames: {str(e)}", "danger")
        print(f"Error: {str(e)}")  # Log the error for debugging
        return redirect(url_for('admin_dashboard'))  # Redirect back to admin dashb


@app.route('/update_role', methods=['POST'])
def update_role():
    if request.method == 'POST':
        new_role = request.form.get('new_role')

        if not new_role:
            flash("Please enter a valid role.", category='danger')
            return redirect(url_for('admin_dashboard'))

        set_current_role(new_role)

        flash(f"Role has been updated to '{new_role}' successfully!", category='success')
        return redirect(url_for('admin_dashboard'))
def generate_users():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))  # Ensure only admin can access this route

    num_voters = request.form.get('num_voters')  # Get the number of voters from the form

    if not num_voters or not num_voters.isdigit() or int(num_voters) <= 0:
        flash("Please enter a valid number of voters.", "danger")
        return redirect(url_for('admin_dashboard'))  # Redirect to admin dashboard if invalid number

    num_voters = int(num_voters)

    # Generate random 4-digit numbers for voters
    generated_usernames = [str(random.randint(1000, 9999)) for _ in range(num_voters)]

    try:
        # Read existing usernames from the 'usernames.txt' file to avoid duplicates
        with open(usernames_file, 'r') as file:
            existing_usernames = set(file.read().splitlines())  # Use a set to avoid duplicates

        # Filter out usernames that already exist in 'usernames.txt'
        new_usernames = [username for username in generated_usernames if username not in existing_usernames]

        if not new_usernames:
            flash("All generated usernames already exist. No new usernames were added.", "warning")
            return redirect(url_for('admin_dashboard'))  # Redirect if no new usernames were generated

        # Update the 'usernames.txt' file with the new unique usernames
        with open(usernames_file, 'a') as file:
            for username in new_usernames:
                file.write(f"{username}\n")

        # Now, generate the Word document with the updated list of usernames
        # Combine the old and new usernames
        all_usernames = list(existing_usernames) + new_usernames
        create_word_document(all_usernames)

        flash(f"{len(new_usernames)} new voters have been added and a Word document has been generated.", "success")
        return redirect(url_for('admin_dashboard'))  # Redirect back to admin dashboard

    except Exception as e:
        flash(f"An error occurred while updating usernames: {str(e)}", "danger")
        print(f"Error: {str(e)}")  # Log the error for debugging
        return redirect(url_for('admin_dashboard'))  # Redirect back to admin dashboard if an error occurs


def create_word_document(usernames, output_file='static/voters_list.docx'):
    doc = Document()
    doc.add_heading('Voters List', level=1)

    for username in usernames:
        doc.add_paragraph(username)

    doc.save(output_file)


def generate_pdf(usernames, output_file='static/voters_list.pdf'):
    c = canvas.Canvas(output_file, pagesize=letter)
    width, height = letter

    # Set up margins and font size for labels
    margin = 50
    column_width = (width - 3 * margin) / 2  # Divide the page into two columns
    label_height = 50  # Increased height of each label for taller cells
    font_size = 12
    max_users_per_column = 10  # Adjusted to fit more height for each user, reduced the max users per column

    c.setFont("Helvetica", font_size)

    # Split the usernames into two groups (columns)
    left_column_users = usernames[:max_users_per_column]
    right_column_users = usernames[max_users_per_column:max_users_per_column*2]

    # Initial y-positions for two columns
    y_position_left = height - margin - label_height  # Starting position for the left column
    y_position_right = height - margin - label_height  # Starting position for the right column

    # Draw the usernames in the left column
    for i, username in enumerate(left_column_users):
        c.setStrokeColor(colors.black)
        c.setFillColor(colors.white)
        c.rect(margin, y_position_left - label_height, column_width, label_height, fill=1)
        c.setFillColor(colors.black)
        c.drawString(margin + 10, y_position_left - label_height + 15, username)
        y_position_left -= label_height + 5  # Move down for the next label

    # Draw the usernames in the right column
    for i, username in enumerate(right_column_users):
        c.setStrokeColor(colors.black)
        c.setFillColor(colors.white)
        c.rect(margin + column_width + margin, y_position_right - label_height, column_width, label_height, fill=1)
        c.setFillColor(colors.black)
        c.drawString(margin + column_width + margin + 10, y_position_right - label_height + 15, username)
        y_position_right -= label_height + 5  # Move down for the next label

    # If there's space left in either column, continue filling in the next page
    if y_position_left <= margin and y_position_right <= margin:
        c.showPage()
        y_position_left = height - margin - label_height
        y_position_right = height - margin - label_height

    c.save()

@app.route('/download_voters_pdf')
def download_voters_pdf():
    try:
        # Path to the generated voters PDF
        file_path = os.path.join('static', 'voters_list.pdf')

        # Check if the file exists
        if not os.path.exists(file_path):
            flash("Voters list PDF not found.", "danger")
            return redirect(url_for('admin_dashboard'))  # Redirect back to the dashboard if the file doesn't exist

        # Serve the file for download
        return send_file(file_path, as_attachment=True)

    except Exception as e:
        flash(f"An error occurred while downloading the file: {str(e)}", "danger")
        return redirect(url_for('admin_dashboard'))  # Redirect back to admin dashboard if an error occurs


@app.route('/view_usernames')
def view_usernames():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))  # Redirect if the user is not admin

    try:
        # Read the usernames from the 'usernames.txt' file
        with open(usernames_file, 'r') as file:
            usernames = file.read().splitlines()  # Read all usernames line by line

        if not usernames:
            flash("No usernames found.", "warning")  # Show a message if no users are found

        return render_template('view_usernames.html', usernames=usernames)  # Pass usernames to template

    except FileNotFoundError:
        flash("Usernames file not found.", "danger")
        return redirect(url_for('admin_dashboard'))  # Redirect back to admin dashboard if the file is missing
@app.route('/delete_all_voters', methods=['POST'])
def delete_all_voters():
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))  # Ensure only admin can access this route

    try:
        # Clear the usernames list by overwriting the file with an empty list
        with open(usernames_file, 'w') as file:
            file.write("")  # Overwrite the file with nothing

        flash("All voters have been deleted.", "success")
    except FileNotFoundError:
        flash("Usernames file not found.", "danger")
    except Exception as e:
        flash(f"An error occurred: {str(e)}", "danger")

    return redirect(url_for('view_usernames'))  # Redirect back to the view_usernames
@app.route('/delete_voter/<username>', methods=['POST'])
def delete_voter(username):
    if 'username' not in session or session['username'] != 'admin':
        flash("Access restricted to admin only.", "danger")
        return redirect(url_for('index'))  # Ensure only admin can access this route

    try:
        # Read the current usernames from the 'usernames.txt' file
        with open(usernames_file, 'r') as file:
            usernames = file.read().splitlines()

        # If the username exists, remove it from the list
        if username in usernames:
            usernames.remove(username)

            # Save the updated usernames back to the file
            with open(usernames_file, 'w') as file:
                for username in usernames:
                    file.write(f"{username}\n")

            flash(f"Voter {username} has been removed.", "success")
        else:
            flash(f"Voter {username} not found.", "warning")

    except FileNotFoundError:
        flash("Usernames file not found.", "danger")
    except Exception as e:
        flash(f"An error occurred: {str(e)}", "danger")

    return redirect(url_for('view_usernames'))  # Redirect back to the view_usernames page

@app.route('/logout')
def logout():
    session.clear()  # Clear all session data
    flash("You have been logged out.", "success")
    return redirect(url_for('index'))  # Redirect to the login page

if __name__ == "__main__":
    app.run(host='0.0.0.0', port=5000, ssl_context=(ssl_cert, ssl_key))
